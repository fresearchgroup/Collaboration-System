from django.shortcuts import render, redirect
from django.urls import reverse_lazy

# Create your views here.
from django.http import Http404, HttpResponse, JsonResponse, HttpResponseRedirect
from .models import Community, CommunityMembership, CommunityArticles, Locations, RequestCommunityCreation, RequestCommunityCreationAssignee, RequestCommunityCreationDetails, CommunityCourses, CommunityMedia, MergedArticles, MergedArticleStates
from rest_framework import viewsets
# from .models import CommunityGroups
from UserRolesPermission.views import user_dashboard
from django.contrib.auth.models import Group as Roles
from rolepermissions.roles import assign_role
from UserRolesPermission.roles import CommunityAdmin
from django.contrib.auth.models import User
from workflow.models import States
from django.db.models import Q
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from Course.views import course_view, create_course
# from notifications.signals import notify
from actstream import action
from actstream.models import Action
from actstream.models import target_stream
from django.contrib.contenttypes.models import ContentType
from feeds.views import update_role_feed,remove_or_add_user_feed
# from notification.views import notify_subscribe_unsubscribe, notify_update_role, notify_remove_or_add_user
from django.conf import settings
from ast import literal_eval
import json
import datetime
import requests
from etherpad.views import create_community_ether, create_article_ether_community, create_session_community
from django.views.generic import CreateView, UpdateView
from .forms import CommunityCreateForm, RequestCommunityCreateForm, CommunityUpdateForm, SubCommunityCreateForm, DocumentForm
from django.contrib import messages
from django.db import connection
from django.urls import reverse
from Categories.models import Category
from PIL import Image
from django.views.generic import TemplateView
from django.views.generic.detail import DetailView
from django.http import JsonResponse
from haystack.generic_views import FacetedSearchView as BaseFacetedSearchView
from haystack.query import SearchQuerySet
from .forms import FacetedProductSearchForm
from UserRolesPermission.models import ProfileImage
from django.db.models import Count
from django.db.models import F
from BasicArticle.models import ArticleStates, Articles
from Media.models import Media, MediaStates
from django.contrib.auth.models import User
import ast
from webcontent.views import mobileBrowser, sendEmail_contributor_content_curated, sendEmail_curator_new_curator_contributions, sendEmail_merged_content_curated, sendEmail_contributor_pow_request_submitted, sendEmail_curator_pow_request_submitted, sendEmail_curate_new_pow
from django.contrib.sites.models import Site
from docx import Document
from htmldocx import HtmlToDocx
import os
from django.template import Context, loader

def display_communities(request):
	communities=Community.objects.filter(parent=None).order_by('name')
	if mobileBrowser(request):
		return render(request, 'communities_mobile.html', {'communities':communities})
	return render(request, 'communities.html', {'communities':communities})

# def display_communities(request):
# 	sorting_by = ["A to Z", "Z to A"]
# 	if request.method == 'POST':
# 		if 'sortingby' in request.POST:
# 			sortselected = request.POST['sortingby']
# 			if sortselected == 'A to Z':
# 				communities=Community.objects.filter(parent=None).order_by('name')
# 			if sortselected == 'Z to A':
# 				communities=Community.objects.filter(parent=None).order_by('-name')
# 			if sortselected == 'oldest':
# 				communities=Community.objects.filter(parent=None).order_by('created_at')
# 			if sortselected == 'latest':
# 				communities=Community.objects.filter(parent=None).order_by('-created_at')
# 		elif 'category' in request.POST:
# 			category = request.POST['category']
# 			category = Category.objects.get(pk=category)
# 			communities=Community.objects.filter(category=category)
# 	else:
# 		sortselected = None
# 		communities=Community.objects.filter(parent=None).order_by('name')
# 	return render(request, 'communities.html',{'communities':communities, 'sorting_by':sorting_by, 'sortselected':sortselected})

def community_view(request, pk):
	try:
		message = 0
		community = Community.objects.get(pk=pk)
		uid = request.user.id
		membership = CommunityMembership.objects.filter(user =uid, community = community.pk)
		role = Roles.objects.get(name='community_admin')
		# if membership.role == role:
		# 	count = CommunityMembership.objects.filter(community=community,role=role).count()
		# 	if count < 2:
		# 		message = 1
	except CommunityMembership.DoesNotExist:
		membership = 'FALSE'
	subscribers = CommunityMembership.objects.filter(community = pk).count()
	pubarticles = CommunityArticles.objects.raw('select ba.id, ba.body, ba.title, workflow_states.name as state from  workflow_states, BasicArticle_articles as ba , Community_communityarticles as ca  where ba.state_id=workflow_states.id and  ca.article_id =ba.id and ca.community_id=%s and ba.state_id in (select id from workflow_states as w where w.name = "publish");', [community.pk])
	#pubarticlescount = len(list(pubarticles))

	top_contributors = CommunityArticles.objects.values('user__username').annotate(num=Count('user__username')).filter(community=community).order_by('-num')[:8]
	for top in top_contributors:
		username = top['user__username']
		try:
			user_profile = ProfileImage.objects.get(user__username=username)
			top['photo'] = user_profile.photo
		except ProfileImage.DoesNotExist:
			top['photo'] = ''

	communitymem=CommunityMembership.objects.filter(community = pk).order_by('?')[:10]
	for comm in communitymem:
		try:
			user_profile = ProfileImage.objects.get(user=comm.user)
			comm.photo = user_profile.photo
		except ProfileImage.DoesNotExist:
			user_profile = "No Image available"
	children = community.get_children().order_by('-contribution_status')

	for child in children:
		curatorrole = Roles.objects.get(name='curator')
		curatorname = CommunityMembership.objects.filter(community=child, role=curatorrole).order_by('-assignedon')[:1]
		if curatorname.count() > 0:
			child.curatorname = curatorname[0].user

		grandchildren = child.get_children()
		draftCount = 0
		submittedCount = 0
		inreviewCount = 0
		modifyCount = 0
		acceptedCount = 0
		rejectedCount = 0
		for gc in grandchildren:
			draftCount += CommunityArticles.objects.filter(community=gc, article__state__name='draft').count()
			draftCount += CommunityMedia.objects.filter(community=gc, media__state__name='draft').count()
			submittedCount += CommunityArticles.objects.filter(community=gc, article__state__name='submitted').count()
			submittedCount += CommunityMedia.objects.filter(community=gc, media__state__name='submitted').count()
			inreviewCount += CommunityArticles.objects.filter(community=gc, article__state__name='reviewStarted').count()
			inreviewCount += CommunityMedia.objects.filter(community=gc, media__state__name='reviewStarted').count()
			modifyCount += CommunityArticles.objects.filter(community=gc, article__state__name='sentToModify').count()
			modifyCount += CommunityMedia.objects.filter(community=gc, media__state__name='sentToModify').count()
			acceptedCount += CommunityArticles.objects.filter(community=gc, article__state__name='accepted').count()
			acceptedCount += CommunityMedia.objects.filter(community=gc, media__state__name='accepted').count()
			acceptedCount += CommunityArticles.objects.filter(community=gc, article__state__name='publish').count()
			acceptedCount += CommunityMedia.objects.filter(community=gc, media__state__name='publish').count()
			acceptedCount += CommunityArticles.objects.filter(community=gc, article__state__name='publishedICP').count()
			acceptedCount += CommunityMedia.objects.filter(community=gc, media__state__name='publishedICP').count()
			rejectedCount += CommunityArticles.objects.filter(community=gc, article__state__name='rejected').count()
			rejectedCount += CommunityMedia.objects.filter(community=gc, media__state__name='rejected').count()

		child.draftCount = draftCount
		child.submittedCount = submittedCount
		child.inreviewCount = inreviewCount
		child.modifyCount = modifyCount
		child.acceptedCount = acceptedCount
		child.rejectedCount = rejectedCount
		if MergedArticles.objects.filter(community=child).exists():
			child.merged = True
			mergedstatus = MergedArticles.objects.filter(community=child)
			child.mergedstatus = mergedstatus[0].state.name
			child.publishedlink = mergedstatus[0].publishedlink
		else:
			child.merged = False

	childrencount = children.count()

	createpow = False
	isCurator = False
	isApprover = False
	if request.user.is_authenticated:
		u = User.objects.get(username=request.user)
		if u.groups.filter(name='curator').exists():
			createpow = True
			isCurator = True
		if u.groups.filter(name='icpapprover').exists():
			isApprover = True

	# if mobileBrowser(request):
	# 	t = loader.get_template('communityview_mobile.html')
	# else:
	# 	t = loader.get_template('communityview.html')
	# c = {'community': community, 'membership':membership, 'subscribers':subscribers, 'top_contributors':top_contributors, 'message':message, 'pubarticles':pubarticles, 'communitymem':communitymem, 'children':children, 'childrencount':childrencount, 'createpow':createpow, 'isCurator':isCurator, 'isApprover':isApprover}
	# return HttpResponse(t.render(c))
	context = {'community': community, 'membership':membership, 'subscribers':subscribers, 'top_contributors':top_contributors, 'message':message, 'pubarticles':pubarticles, 'communitymem':communitymem, 'children':children, 'childrencount':childrencount, 'createpow':createpow, 'isCurator':isCurator, 'isApprover':isApprover}
	if mobileBrowser(request):
		return render(request, 'communityview_mobile.html', context)
	return render(request, 'communityview.html', context)

def community_subscribe(request):
	cid = request.POST['cid']
	if request.user.is_authenticated:
		if request.method == 'POST':
			community=Community.objects.get(pk=cid)
			role = Roles.objects.get(name='author')
			user = request.user

			if CommunityMembership.objects.filter(user=user, community=community).exists():
				return redirect('community_view',pk=cid)
			# notify_subscribe_unsubscribe(request.user,community, 'subscribe')
			obj = CommunityMembership.objects.create(user=user, community=community, role=role)
			return redirect('community_view',pk=cid)
		return render(request, 'communityview.html')
	else:
		return redirect('/login/?next=/community-view/%d' % int(cid) )



def community_unsubscribe(request):
	if request.user.is_authenticated:
		if request.method == 'POST':
			cid = request.POST['cid']
			community=Community.objects.get(pk=cid)
			user = request.user
			if CommunityMembership.objects.filter(user=user, community=community).exists():
				remove_or_add_user_feed(user,community,'left')
				# notify_remove_or_add_user(user, user, community, 'left')
				obj = CommunityMembership.objects.filter(user=user, community=community).delete()
				#notify_subscribe_unsubscribe(request.user, community, 'unsubscribe')
			return redirect('community_view',pk=cid)
		return render(request, 'communityview.html')
	else:
		return redirect('login')

class RequestCommunityCreationView(CreateView):
	form_class = RequestCommunityCreateForm
	model = RequestCommunityCreationDetails
	template_name = 'request_community_creation.html'
	success_url = 'user_dashboard'
	context_object_name = 'context'

	def get_context_data(self, **kwargs):
		context = super().get_context_data(**kwargs)
		community = Community.objects.filter(parent__pk=self.request.GET.get('cid'))
		context['community'] = community
		return context

	def get_form_kwargs(self):
		kwargs = super(RequestCommunityCreationView, self).get_form_kwargs()
		kwargs['cid'] = self.request.GET.get('cid')
		return kwargs

	def get(self, request, *args, **kwargs):
		if request.user.is_authenticated:
			self.object = None
			return super(RequestCommunityCreationView, self).get(request, *args, **kwargs)
		return redirect('login')

	def form_valid(self, form):
		"""
		If the form is valid, save the associated model.
		"""
		self.object = form.save(commit=False)
		self.object.actionby = self.request.user

		self.object.state = form.cleaned_data.get('state')
		if form.cleaned_data.get('district') == "Others":
			self.object.district = self.request.POST['districtOthers']
		else:
			self.object.district = form.cleaned_data.get('district')

		if form.cleaned_data.get('city') == "Others":
			self.object.city = self.request.POST['cityOthers']
		else:
			self.object.city = form.cleaned_data.get('city')

		if form.cleaned_data.get('pincode') == "Others":
			self.object.pincode = self.request.POST['pincodeOthers']
		else:
			self.object.pincode = form.cleaned_data.get('pincode')
		
		cid = form.cleaned_data.get('parent')
		parent = Community.objects.get(name=cid)

		rcommunity = RequestCommunityCreation.objects.create(
			requestedon = datetime.datetime.now(),
			requestedby = self.request.user,
			email = self.request.user.email,
			parent = parent
		)
		self.object.requestcommunity = rcommunity
		self.object.save()

		curators = User.objects.filter(groups__name='curator').exclude(username='admin').order_by('pk')
		curatorCount = curators.count()
		requestCount = RequestCommunityCreation.objects.all().count()
		curatorNo = requestCount % curatorCount
		assignedto = curators[curatorNo]

		RequestCommunityCreationAssignee.objects.create(
			requestcommunity = rcommunity,
			assignedto = assignedto,
			assignedon = datetime.datetime.now()
		)

		to = []
		to.append(self.request.user.email)
		sendEmail_contributor_pow_request_submitted(to)
		to = []
		to.append(assignedto.email)
		sendEmail_curator_pow_request_submitted(to)

		messages.success(self.request, 'Request for creation of place of worship successfully submited.')
		return super(RequestCommunityCreationView, self).form_valid(form)

	def get_success_url(self):
		"""
		Returns the supplied URL.
		"""
		if self.success_url:
			return reverse(self.success_url)
		else:
			try:
				url = self.object.get_absolute_url()
			except AttributeError:
				raise ImproperlyConfigured(
				"No URL to redirect to.  Either provide a url or define"
				" a get_absolute_url method on the Model.")
		return url

def update_community_requests(request):
	pk = request.POST['pk']
	rcommunity=RequestCommunityCreation.objects.get(pk=pk)
	user = request.user
	name = request.POST['name']
	desc = request.POST['description']
	area = request.POST['area']
	city = request.POST['city']
	district = request.POST['district']
	latitude = request.POST['latitude']
	longitude = request.POST['longitude']
	state = request.POST['state']
	pincode = request.POST['pincode']
	reason = request.POST['reason']

	RequestCommunityCreationDetails.objects.create(
		requestcommunity = rcommunity,
		name = name,
		desc = desc,
		area = area,
		city = city,
		district = district,
		latitude = latitude,
		longitude = longitude,
		state = state,
		pincode = pincode,
		status = 'Requested',
		actionby = user,
		reason = reason,
		actionon = datetime.datetime.now()
	)

	to = []
	to.append(request.user.email)
	sendEmail_contributor_pow_request_submitted(to)
	to = []
	uname = RequestCommunityCreationAssignee.objects.filter(requestcommunity__id=pk).order_by('-assignedon')[:1]
	to.append(uname[0].assignedto.email)
	sendEmail_curator_pow_request_submitted(to)
	return redirect('user_dashboard')

def handle_community_creation_requests(request):

	u = User.objects.get(username=request.user)
	if u.groups.filter(name='curator').exists():
		# if request.user.is_superuser:
		if request.method == 'POST':
			pk = request.POST['pk']
			rcommunity=RequestCommunityCreation.objects.get(pk=pk)
			user = request.user
			community_parent = request.POST['community_parent']
			parent = Community.objects.get(pk=community_parent)
			status = request.POST['status']
			reason = ''
			pow = RequestCommunityCreationDetails.objects.filter(requestcommunity__id=pk).order_by('-actionon')[:1]
			pow = pow[0].name
			uname = ''
			to = []

			if status == 'changeassignee':
				uname = RequestCommunityCreationAssignee.objects.filter(requestcommunity__id=pk).order_by('-assignedon')[:1]
				to.append(uname[0].assignedto.email)
				RequestCommunityCreationAssignee.objects.create(
					requestcommunity = rcommunity,
					assignedto = user,
					assignedon = datetime.datetime.now()
				)
				uname = request.user.username

			if status=='accept':
				name = request.POST['name']
				desc = request.POST['description']
				area = request.POST['area']
				city = request.POST['city']
				state = request.POST['state']
				district = request.POST['district']
				pincode = request.POST['pincode']
				latitude = request.POST['latitude']
				longitude = request.POST['longitude']
				communitycreation = Community.objects.create(
					name = name,
					desc = desc,
					area = area,
					city = city,
					state = state,
					pincode = pincode,
					district = district,
					latitude = latitude,
					longitude = longitude,
					created_by = user,
					parent = parent,
					contributed_by = rcommunity.requestedby
				)

				RequestCommunityCreationDetails.objects.create(
					requestcommunity = rcommunity,
					name = communitycreation.name,
					desc = communitycreation.desc,
					area = communitycreation.area,
					city = communitycreation.city,
					state = communitycreation.state,
					pincode = communitycreation.pincode,
					district = communitycreation.district,
					latitude = communitycreation.latitude,
					longitude = communitycreation.longitude,
					status = 'accepted',
					actionby = user,
					actionon = datetime.datetime.now()
				)
				Community.objects.create(name='Introduction', created_by = user, parent = communitycreation)
				Community.objects.create(name='Architecture', created_by = user, parent = communitycreation)
				Community.objects.create(name='Rituals', created_by = user, parent = communitycreation)
				Community.objects.create(name='Ceremonies', created_by = user, parent = communitycreation)
				Community.objects.create(name='Tales', created_by = user, parent = communitycreation)
				Community.objects.create(name='More Information', created_by = user, parent = communitycreation)
				
				communityadmin = Roles.objects.get(name='community_admin')
				CommunityMembership.objects.create(
					user = user,
					community = communitycreation,
					role = communityadmin
				)
				to.append(rcommunity.requestedby.email)
				remove_or_add_user_feed(rcommunity.requestedby,communitycreation,'community_created')
		
			if status=='modify' or status=='rejected':
				reason = request.POST['reason']
				rcom = RequestCommunityCreationDetails.objects.filter(requestcommunity__id=pk).order_by('-actionon')[:1]
				RequestCommunityCreationDetails.objects.create(
					requestcommunity = rcommunity,
					name = rcom[0].name,
					desc = rcom[0].desc,
					area = rcom[0].area,
					city = rcom[0].city,
					state = rcom[0].state,
					pincode = rcom[0].pincode,
					district = rcom[0].district,
					latitude = rcom[0].latitude,
					longitude = rcom[0].longitude,
					status = status,
					reason = reason,
					actionby = user,
					actionon = datetime.datetime.now()
				)
				to.append(rcommunity.requestedby.email)

			sendEmail_curate_new_pow(to, pow, parent.name, reason, uname, status)

		rids = RequestCommunityCreation.objects.all().values_list('pk', flat=True)
		requestcommunitycreation = []
		for i in rids:
			rcom = RequestCommunityCreationDetails.objects.filter(requestcommunity__id=i).order_by('-actionon')[:1]
			assignees = RequestCommunityCreationAssignee.objects.filter(requestcommunity__id=i).order_by('-assignedon')[:1]
			for r in rcom:
				r.assignedto = assignees[0].assignedto
				r.assignedon = assignees[0].assignedon
				requestcommunitycreation.append(r)
		return render(request, 'community_creation_requests.html',{'requestcommunitycreation':requestcommunitycreation})
	else:
		return redirect('login')

def manage_community(request,pk):
	if request.user.is_authenticated:
		community = Community.objects.get(pk=pk)
		uid = request.user.id
		errormessage = ''
		membership = None
		try:
			membership = CommunityMembership.objects.get(user =uid, community = community.pk)
			if membership.role.name == 'community_admin':
				count = CommunityMembership.objects.filter(community = community.pk, role=membership.role).count()
				members = CommunityMembership.objects.filter(community = community.pk)
				if request.method == 'POST':
					try:
						username = request.POST['username']
						rolename = request.POST['role']
						user = User.objects.get(username = username)
						role = Roles.objects.get(name=rolename)
						status = request.POST['status']

						if status == 'add':
							try:
								is_member = CommunityMembership.objects.get(user =user, community = community.pk)
							except CommunityMembership.DoesNotExist:
								obj = CommunityMembership.objects.create(user=user, community=community, role=role)
								# notify_remove_or_add_user(request.user, user, community, 'added')
								if rolename=='publisher' or rolename=='community_admin':
									remove_or_add_user_feed(user,community,'added')

							else:
								errormessage = 'user exists in community'
						if status == 'update':
							if count > 1 or count == 1 and username != request.user.username:
								try:
									update_role_feed(user,community,rolename)
									# notify_update_role(request.user, user,community,rolename)
									is_member = CommunityMembership.objects.get(user =user, community = community.pk)
									is_member.role = role
									is_member.save()


								except CommunityMembership.DoesNotExist:
									errormessage = 'no such user in the community'
							else:
								errormessage = 'cannot update this user'
						if status == 'remove':
							if count > 1 or count == 1 and username != request.user.username:
								try:
									remove_or_add_user_feed(user,community,'removed')
									# notify_remove_or_add_user(request.user, user,community,'removed')
									obj = CommunityMembership.objects.filter(user=user, community=community).delete()



								except CommunityMembership.DoesNotExist:
									errormessage = 'no such user in the community'
							else:
								errormessage = 'cannot remove this user'
						return render(request, 'managecommunity.html', {'community': community, 'members':members,'membership':membership, 'errormessage':errormessage})
	#					return redirect('manage_community',pk=pk)
					except User.DoesNotExist:
						errormessage = "no such user in the system"

				return render(request, 'managecommunity.html', {'community': community, 'members':members,'membership':membership, 'errormessage':errormessage})
			else:
				return redirect('community_view',pk=pk)
		except CommunityMembership.DoesNotExist:
			return redirect('community_view',pk=pk)
	else:
		return redirect('login')

class UpdateCommunityView(UpdateView):
	form_class = CommunityUpdateForm
	model = Community
	template_name = 'updatecommunityinfo.html'
	#community_admin = Roles.objects.get(name='community_admin')
	success_url = 'community_view'
	pk_url_kwarg = 'pk'
	context_object_name = 'community'

	def get(self, request, *args, **kwargs):
		if request.user.is_authenticated:
			u = User.objects.get(username=request.user)
			if u.groups.filter(name='curator').exists():
				self.object = self.get_object()
				return super(UpdateCommunityView, self).get(request, *args, **kwargs)
			messages.warning(self.request, 'You cannot update a pow')
			return redirect(self.success_url, self.object.pk)
		return redirect('home')

	def get_context_data(self, **kwargs):
		context = super().get_context_data(**kwargs)
		return context

	def get_membership(self):
		if CommunityMembership.objects.filter(user =self.request.user, community = self.object).exists():
			return CommunityMembership.objects.get(user =self.request.user, community = self.object)
		return False

	def form_valid(self, form):
		"""
		If the form is valid, save the associated model.
		"""
		self.object = form.save(commit=False)
		# self.object.image_thumbnail = form.cleaned_data.get('image')

		self.object.state = form.cleaned_data.get('state')
		if form.cleaned_data.get('district') == "Others":
			self.object.district = self.request.POST['districtOthers']
		else:
			self.object.district = form.cleaned_data.get('district')

		if form.cleaned_data.get('city') == "Others":
			self.object.city = self.request.POST['cityOthers']
		else:
			self.object.city = form.cleaned_data.get('city')

		if form.cleaned_data.get('pincode') == "Others":
			self.object.pincode = self.request.POST['pincodeOthers']
		else:
			self.object.pincode = form.cleaned_data.get('pincode')

		self.object.save()
		
		# if form.cleaned_data.get('x'):
		# 	x = form.cleaned_data.get('x')
		# 	y = form.cleaned_data.get('y')
		# 	w = form.cleaned_data.get('width')
		# 	h = form.cleaned_data.get('height')
		# 	image = Image.open(self.object.image_thumbnail)
		# 	cropped_image = image.crop((x, y, w+x, h+y))
		# 	resized_image = cropped_image.resize((200, 200), Image.ANTIALIAS)
		# 	resized_image.save(self.object.image_thumbnail.path)

		return super(UpdateCommunityView, self).form_valid(form)

	def get_success_url(self):
		"""
		Returns the supplied URL.
		"""
		if self.success_url:
			return reverse(self.success_url,kwargs={'pk': self.object.parent.pk})
		else:
			try:
				url = self.object.get_absolute_url()
			except AttributeError:
				raise ImproperlyConfigured(
				"No URL to redirect to.  Either provide a url or define"
				" a get_absolute_url method on the Model.")
		return url


class CreateCommunityView(CreateView):
	form_class = CommunityCreateForm
	model = Community
	template_name = 'create_community.html'
	#community_admin = Roles.objects.get(name='community_admin')
	success_url = 'community_view'

	def get(self, request, *args, **kwargs):
		if request.user.is_superuser:
			self.object = None
			return super(CreateCommunityView, self).get(request, *args, **kwargs)
		return redirect('home')

	def form_valid(self, form):
		"""
		If the form is valid, save the associated model.
		"""
		self.object = form.save(commit=False)
		self.object.image_thumbnail = form.cleaned_data.get('image')

		# forum_link, fid = self.create_forum(self.object.name, self.object.desc)
		# if forum_link is not False:
		# 	self.object.forum_link = forum_link
		# 	self.object.forum = fid
		# else:
		# 	messages.warning(self.request, 'Cannot Create Forum for this community. Please check if default forum is created.')
		# 	return super(CreateCommunityView, self).form_invalid(form)
		self.object.save()

		if self.object.image_thumbnail:
			x = form.cleaned_data.get('x')
			y = form.cleaned_data.get('y')
			w = form.cleaned_data.get('width')
			h = form.cleaned_data.get('height')
			image = Image.open(self.object.image_thumbnail)
			cropped_image = image.crop((x, y, w+x, h+y))
			resized_image = cropped_image.resize((200, 200), Image.ANTIALIAS)
			resized_image.save(self.object.image_thumbnail.path)

		CommunityMembership.objects.create(
			user = self.object.created_by,
			community = self.object,
			role = Roles.objects.get(name='community_admin')
			)

		#create the ether id for community
		try:
			create_community_ether(self.object)
		except Exception as e:
			messages.warning(self.request, 'Cannot create ether id for this Community. Please check whether Etherpad service is running.')

		#create the ether id for community
		try:
			create_wiki_for_community(self.object)
		except Exception as e:
			messages.warning(self.request, 'Cannot create wiki for this Community. Please check default wiki is created.')

		return super(CreateCommunityView, self).form_valid(form)

	def create_forum(self, name, desc):
		# Create Forum for this community
		try:
			cursor = connection.cursor()
			cursor.execute(''' select tree_id from forum_forum order by tree_id DESC limit 1''')
			try:
				tree_id = cursor.fetchone()[0] + 1
			except:
				tree_id = 0
			slug = "-".join(name.lower().split())
			#return HttpResponse(str(tree_id))
			insert_stmt = (
				  "INSERT INTO forum_forum (created,updated,name,slug,description,link_redirects,type,link_redirects_count,display_sub_forum_list,lft,rght,tree_id,level,direct_posts_count,direct_topics_count) "
				  "VALUES (NOW(), NOW(), %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"
				)
			data = (name, slug, desc, 0,0,0,1,1,2,tree_id,0,0,0)
			cursor.execute(insert_stmt, data)
			cursor.execute(''' select id from forum_forum order by id desc limit 1''')
			fid = str(cursor.fetchone()[0])
			forum_link = slug + '-' + fid
			return forum_link, fid
		except:
			return False, False

	def get_success_url(self):
		"""
		Returns the supplied URL.
		"""
		if self.success_url:
			return reverse(self.success_url,kwargs={'pk': self.object.pk})
		else:
			try:
				url = self.object.get_absolute_url()
			except AttributeError:
				raise ImproperlyConfigured(
				"No URL to redirect to.  Either provide a url or define"
				" a get_absolute_url method on the Model.")
		return url

class CreateSubCommunityView(CreateView):
	form_class = SubCommunityCreateForm
	model = Community
	template_name = 'create_community.html'
	success_url = 'community_view'

	def get_context_data(self, **kwargs):
		context = super().get_context_data(**kwargs)
		community = Community.objects.filter(parent__pk=self.kwargs['pk'])
		context['community'] = community
		return context

	def get_initial(self):
		"""
		Returns the initial data to use for forms on this view.
		"""
		initial= self.initial.copy()
		initial.update({'parent': self.model.objects.get(pk=self.kwargs['pk']) })
		return initial

	def get(self, request, *args, **kwargs):
		if request.user.is_authenticated:
			u = User.objects.get(username=request.user)
			if u.groups.filter(name='curator').exists():
				self.object = None
				return super(CreateSubCommunityView, self).get(request, *args, **kwargs)
			messages.warning(self.request, 'You cannot create a pow')
			return redirect(self.success_url, self.kwargs['pk'])
		return redirect('home')

	def is_member_of_parent(self):
		community = self.model.objects.get(pk=self.kwargs['pk'])
		return CommunityMembership.objects.filter(community=community, user=self.request.user).exists()

	def form_valid(self, form):
		"""
		If the form is valid, save the associated model.
		"""
		self.object = form.save(commit=False)
		self.object.created_by = self.request.user
		self.object.image_thumbnail = form.cleaned_data.get('image')
		self.object.save()

		if self.object.image_thumbnail:
			x = form.cleaned_data.get('x')
			y = form.cleaned_data.get('y')
			w = form.cleaned_data.get('width')
			h = form.cleaned_data.get('height')
			image = Image.open(self.object.image_thumbnail)
			cropped_image = image.crop((x, y, w+x, h+y))
			resized_image = cropped_image.resize((200, 200), Image.ANTIALIAS)
			resized_image.save(self.object.image_thumbnail.path)

		CommunityMembership.objects.create(
			user = self.object.created_by,
			community = self.object,
			role = Roles.objects.get(name='community_admin')
			)

		Community.objects.create(name='Introduction', created_by = self.object.created_by, parent = self.object)
		Community.objects.create(name='Architecture', created_by = self.object.created_by, parent = self.object)
		Community.objects.create(name='Rituals', created_by = self.object.created_by, parent = self.object)
		Community.objects.create(name='Ceremonies', created_by = self.object.created_by, parent = self.object)
		Community.objects.create(name='Tales', created_by = self.object.created_by, parent = self.object)
		Community.objects.create(name='More Information', created_by = self.object.created_by, parent = self.object)

		# try:
		# 	create_community_ether(self.object)
		# except Exception as e:
		# 	messages.warning(self.request, 'Cannot create ether id for this Community. Please check whether Etherpad service is running.')

		return super(CreateSubCommunityView, self).form_valid(form)

	def get_form_kwargs(self):
		"""
		Returns the keyword arguments for instantiating the form.
		"""
		kwargs = super(CreateSubCommunityView, self).get_form_kwargs()
		kwargs.update({'community': self.model.objects.get(pk=self.kwargs['pk'])})
		return kwargs

	def get_success_url(self):
		"""
		Returns the supplied URL.
		"""
		if self.success_url:
			return reverse(self.success_url,kwargs={'pk': self.object.pk})
		else:
			try:
				url = self.object.get_absolute_url()
			except AttributeError:
				raise ImproperlyConfigured(
				"No URL to redirect to.  Either provide a url or define"
				" a get_absolute_url method on the Model.")
		return url


def community_content(request, pk):
	commarticles = ''
	try:
		community = Community.objects.get(pk=pk)
		uid = request.user.id
		membership = CommunityMembership.objects.get(user=uid, community=community.pk)
		if membership:

			carticles = CommunityArticles.objects.filter(community=community, article__state__initial=False, article__state__final=False).annotate(title=F('article__title'),state=F('article__state__name'),created_at=F('article__created_at'),body=F('article__body'),image=F('article__image'),views=F('article__views'),username=F('article__created_by__username'))
			for art in carticles:
				art.type = 'article'

			ccourse = CommunityCourses.objects.filter(community=community, course__state__initial=False, course__state__final=False).annotate(title=F('course__title'),state=F('course__state__name'),created_at=F('course__created_at'),body=F('course__body'),image=F('course__image'),username=F('course__created_by__username'))
			for course in ccourse:
				course.type = 'course'

			cmedia = CommunityMedia.objects.filter(community=community, media__state__initial=False, media__state__final=False).annotate(title=F('media__title'),state=F('media__state__name'),created_at=F('media__created_at'),mediatype=F('media__mediatype'),image=F('media__mediafile'),username=F('media__created_by__username'))
			for media in cmedia:
				media.type = 'media'

			ch5p = []
			print('new test')
			try:
				response = requests.get(settings.H5P_ROOT + 'h5p/h5papi/?format=json')
				json_data = json.loads(response.text)
				print(json_data)

				for obj in json_data:
					if obj['community_id'] == community.pk:
						obj['type'] = 'h5p'
						ch5p.append(obj)
			except Exception as e:
				print(e)
				print("H5P server down...Sorry!! We will be back soon")
			lstfinal = list(carticles) +  list(cmedia) + list(ccourse) + list(ch5p)

			page = request.GET.get('page', 1)
			paginator = Paginator(list(lstfinal), 5)
			try:
				commarticles = paginator.page(page)
			except PageNotAnInteger:
				commarticles = paginator.page(1)
			except EmptyPage:
				commarticles = paginator.page(paginator.num_pages)

	except CommunityMembership.DoesNotExist:
		return redirect('community_view', community.pk)
	return render(request, 'communitycontent.html', {'community': community, 'membership':membership, 'commarticles':commarticles})


def h5p_view(request, pk):
	try:
		requests.get(settings.H5P_ROOT + 'h5p/h5papi/?format=json')
		return redirect( settings.H5P_ROOT + "h5p/content/?contentId=%s" % pk)
	except ConnectionError:
		return render(request, 'h5pserverdown', {})

def community_course_create(request):
	if request.user.is_authenticated:
		if request.method == 'POST':
			status = request.POST['status']
			cid = request.POST['cid']
			community = Community.objects.get(pk=cid)
			if status=='1':
				course = create_course(request)
				obj = CommunityCourses.objects.create(course=course, user=request.user, community=community)
				return redirect('course_view', course.pk)
			else:
				return render(request, 'new_course.html', {'community':community, 'status':1})
		else:
			return redirect('home')
	else:
		return redirect('login')



def feed_content(request, pk):
	communityfeed = ''
	try:
		community = Community.objects.get(pk=pk)
		uid = request.user.id
		membership = CommunityMembership.objects.get(user=uid, community=community.pk)
		if membership:
			feeds = community.target_actions.all()
			page = request.GET.get('page', 1)
			paginator = Paginator(feeds, 10)
			try:
				communityfeed = paginator.page(page)
			except PageNotAnInteger:
				communityfeed = paginator.page(1)
			except EmptyPage:
				communityfeed = paginator.page(paginator.num_pages)

	except CommunityMembership.DoesNotExist:
		return redirect('community_view', community.pk)

	return render(request, 'communityfeed.html', {'community': community, 'membership':membership, 'feeds':communityfeed})

def community_h5p_create(request):
	if request.user.is_authenticated:
		if request.method == 'POST':
			cid = request.POST['cid']
			request.session['cid'] = cid
			request.session['gid'] = 0
			try:
				requests.get(settings.H5P_ROOT + 'h5p/h5papi/?format=json')
				return redirect(settings.H5P_ROOT + 'h5p/create/')
			except Exception as e:
				print(e)
				return render(request, 'h5pserverdown.html', {})
		return redirect('home')
	return redirect('login')

def create_wiki_for_community(community):

	cursor = connection.cursor()
	wiki_slug = str(community.name) + str(community.pk)
	#Create wiki for this community
	cursor.execute('''SET FOREIGN_KEY_CHECKS=0''')
	insert_stmt_urlpath = (
			"INSERT INTO wiki_urlpath (id, slug, lft, rght, tree_id, level, article_id, parent_id, site_id)"
			"VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)"
		)

	cursor.execute(''' select rght from wiki_urlpath order by rght DESC limit 1''')
	url_rght = cursor.fetchone()[0]

	cursor.execute(''' select id from wiki_urlpath order by id DESC limit 1''')
	urlpath_id = cursor.fetchone()[0] + 1

	cursor.execute(''' select id from wiki_article order by id DESC limit 1''')
	new_id = cursor.fetchone()[0] + 1

	data_urlpath = (urlpath_id, wiki_slug , url_rght, url_rght + 1, 1, 1, new_id, 1, 1)

	cursor.execute('''update wiki_urlpath set rght = rght + 2 where slug IS NULL''')

	cursor.execute(insert_stmt_urlpath, data_urlpath)

	cursor.execute(''' select id from wiki_articlerevision order by id DESC limit 1''')
	cur_rev_id =  cursor.fetchone()[0] + 1

	insert_stmt_article = (
					"INSERT INTO wiki_article (id, created, modified, group_read, group_write, other_read, other_write,current_revision_id, group_id, owner_id)"
					"VALUES (%s, NOW(), NOW(), %s, %s, %s, %s, %s, %s, %s)"
				)

	data_article = (new_id, 1, 1, 1, 1, cur_rev_id, 1,2)
	cursor.execute(insert_stmt_article, data_article)

	cursor.execute(''' select content_type_id from wiki_articleforobject order by content_type_id DESC limit 1''')
	con_type_id = cursor.fetchone()[0]

	cursor.execute(''' select id from wiki_articleforobject order by id DESC limit 1''')
	forobject_id = cursor.fetchone()[0] + 1

	insert_stmt_articleforobject = (
					"INSERT INTO wiki_articleforobject (id, object_id, is_mptt, article_id, content_type_id)"
					"VALUES (%s, %s, %s, %s, %s)"
				)

	data_articleforobject = (forobject_id, new_id, 1, new_id, con_type_id)

	cursor.execute(insert_stmt_articleforobject, data_articleforobject)

	cursor.execute(''' select id from wiki_articlerevision order by id DESC limit 1''')
	articlerev_id = cursor.fetchone()[0] + 1

	insert_stmt_wikiarticlerevision = (
					"INSERT INTO wiki_articlerevision (id, revision_number, user_message, automatic_log, ip_address, modified, created, deleted, locked, content, title, article_id, user_id)"
					"VALUES (%s, %s, %s, %s, %s, NOW(), NOW(), %s, %s, %s, %s, %s, %s)"
				)

	data_wikiarticlerevision = (articlerev_id ,1, "Write your summary here.", "", "", 0, 0, "Write your content here", str(community.name) , new_id, 2)


	cursor.execute(insert_stmt_wikiarticlerevision, data_wikiarticlerevision)


	cursor.execute('''SET FOREIGN_KEY_CHECKS=1''')

def community_media_create(request):
	if request.user.is_authenticated:
		if request.method == 'POST':
			status = request.POST['status']
			cid = request.POST['cid']
			community = Community.objects.get(pk=cid)
			if status=='1':
				media = create_media(request)
				metadata = create_metadata(request)
				CommunityMedia.objects.create(media=media, user=request.user, community=community)
				MediaMetadata.objects.create(media=media, metadata=metadata)
				return redirect('media_view', media.pk)
			else:
				return render(request, 'new_media.html', {'community':community, 'status':1})
		else:
			return redirect('home')
	else:
		return redirect('login')


class FacetedSearchView(BaseFacetedSearchView, TemplateView):

    form_class = FacetedProductSearchForm
    facet_fields = ['created_at', 'views']
    template_name = 'search_result.html'
    #paginate_by = 3
    # context_object_name = 'object_list'


    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        print("context..................")
        print(context)
        print("self.request.GET >>>>>>>>>>>>>>>>>>>>> ", self.request.GET)
        
        try:
        	if 'community' in self.request.GET:
		        context['community'] = self.request.GET['community']
        	if 'article' in self.request.GET:
	            context['article'] = self.request.GET['article']
        	if 'image' in self.request.GET:
	        	context['image'] = self.request.GET['image']
        	if 'audio' in self.request.GET:
		        context['audio'] = self.request.GET['audio']
        	if 'video' in self.request.GET:
		        context['video'] = self.request.GET['video']
        	if 'query' in self.request.GET:
		        context['query'] = self.request.GET['query']
        except:
            pass
        
        # new_created_at = []
        # for d in context['facets']['fields']['created_at']:
        # 	date_value = datetime.datetime.fromtimestamp(int(str(d[0]))/1000)
        # 	temp_tuple = (date_value, d[1])
        # 	new_created_at.append(temp_tuple)
        # context['facets']['fields']['created_at'] = new_created_at

        object_list = context['object_list']
        return context

def assign_community_curation(request):
	pk = request.POST['community_parent']
	community = Community.objects.get(pk=pk)
	role = Roles.objects.get(name='curator')

	curatorname = CommunityMembership.objects.filter(community=community, role=role).order_by('-assignedon')[:1]
	to = []
	to.append(curatorname[0].user.email)
	newcurator = request.user.username
	sendEmail_curator_new_curator_contributions(to, newcurator, community.name)

	CommunityMembership.objects.create(
		user = request.user,
		community = community,
		role = role,
		assignedon = datetime.datetime.now()
	)
	if 'articlepk' in request.POST:
		articlepk = request.POST['articlepk']
		return redirect('article_view',pk=articlepk)
	elif 'mediapk' in request.POST:
		mediapk = request.POST['mediapk']
		return redirect('media_view',pk=mediapk)
	else:
		return redirect('view_all_content', pk1=pk, state='accepted')

def curate_content(request):
	u = User.objects.get(username=request.user)
	if u.groups.filter(name='curator').exists():
		if request.method == 'POST':
			pk = request.POST['pk']
			conttype = request.POST['type']
			status = request.POST['status']
			comments = ''
			if status == 'markreview':
				state = States.objects.get(name='reviewStarted')
			if status == 'accept':
				state = States.objects.get(name='accepted')
			if status == 'modify':
				state = States.objects.get(name='sentToModify')
				comments = request.POST['reason']
			if status == 'reject':
				state = States.objects.get(name='rejected')
				comments = request.POST['reason']

			if conttype == 'Article':
				article = Articles.objects.get(pk=pk)
				article.state = state
				article.save()
				ArticleStates.objects.create(
					article = article,
					state = state,
					changedby = request.user,
					changedon = datetime.datetime.now(),
					body = article.body,
					comments = comments
				)
				commarticles = article.communityarticles.all()
				section = commarticles[0].community.name
				parent = commarticles[0].community.parent.name
				to = []
				to.append(article.created_by.email)
				sendEmail_contributor_content_curated(to, status, section, parent, comments, request.META.get('HTTP_REFERER'))
				if 'redirecto' in request.POST:
					redirecto = request.POST['redirecto']
					if redirecto == 'view_all_content':
						commpk = request.POST['commpk']
						state =  request.POST['state']
						return redirect('view_all_content', commpk, state)
				return redirect('article_view',pk=pk)

			if conttype == 'Media':
				media = Media.objects.get(pk=pk)
				media.state = state
				media.save()
				MediaStates.objects.create(
					media = media,
					state = state,
					changedby = request.user,
					changedon = datetime.datetime.now(),
					comments = comments
				)
				commmedia = media.communitymedia.all()
				section = commmedia[0].community.name
				parent = commmedia[0].community.parent.name
				to = []
				to.append(media.created_by.email)
				sendEmail_contributor_content_curated(to, status, section, parent, comments, request.META.get('HTTP_REFERER'))
				if 'redirecto' in request.POST:
					redirecto = request.POST['redirecto']
					if redirecto == 'view_all_content':
						commpk = request.POST['commpk']
						state =  request.POST['state']
						return redirect('view_all_content', commpk, state)
				return redirect('media_view',pk=pk)

def display_curation_list(request, pk1='', pk2=''):
	# pk1-community pk, pk2-state name
	if pk1 and pk2:
		community = Community.objects.get(pk=pk1)
		if pk2 == 'accepted':
			states = ['accepted', 'publish', 'publishedICP']
			commarticles = CommunityArticles.objects.filter(community__parent=community, article__state__name__in=states)
			commmedia = CommunityMedia.objects.filter(community__parent=community, media__state__name__in=states)
		else:
			state = States.objects.get(name=pk2)
			commarticles = CommunityArticles.objects.filter(community__parent=community, article__state=state)
			commmedia = CommunityMedia.objects.filter(community__parent=community, media__state=state)
	if not (pk1 and pk2):
		commarticles = CommunityArticles.objects.filter( Q(article__state__name='submitted') |Q(article__state__name='submitted') | Q(article__state__name='reviewStarted') | Q(article__state__name='sentToModify') | Q(article__state__name='accepted') |Q(article__state__name='rejected'))
		commmedia = CommunityMedia.objects.filter( Q(media__state__name='submitted') | Q(media__state__name='reviewStarted') | Q(media__state__name='sentToModify') |Q(media__state__name='accepted') | Q(media__state__name='rejected'))
	lstContent = get_content(commarticles, commmedia)
	return render(request, 'curate_content.html',{'lstContent':lstContent})

def get_content(commarticles, commmedia):
	for cart in commarticles:
		role = Roles.objects.get(name='curator')
		commembership = CommunityMembership.objects.filter(community=cart.community.parent, role=role).order_by('-assignedon')[:1]
		cart.type = 'Article'
		articlestate = ArticleStates.objects.filter(article=cart.article).order_by('-changedon')[:1]
		cart.changedon = articlestate[0].changedon
		cart.changedby = articlestate[0].changedby
		cart.comments = articlestate[0].comments
		cart.assignedto = commembership[0].user.username
		cart.assignedon = commembership[0].assignedon
		currentday = datetime.datetime.now().date() 
		changedon = articlestate[0].changedon.date()
		diff = (currentday - changedon).days
		if cart.article.state.name == 'submitted' or cart.article.state.name == 'reviewStarted' or cart.article.state.name == 'sentToModify':
			if int(diff) > int(settings.DEADLINE):
				cart.deadlinepassed = True
			else:
				cart.deadlinepassed = False
		else:
			cart.deadlinepassed = False
	for cmedia in commmedia:
		role = Roles.objects.get(name='curator')
		commembership = CommunityMembership.objects.filter(community=cmedia.community.parent, role=role).order_by('-assignedon')[:1]
		cmedia.type = 'Media'
		mediastate = MediaStates.objects.filter(media=cmedia.media).order_by('-changedon')[:1]
		cmedia.changedon = mediastate[0].changedon
		cmedia.changedby = mediastate[0].changedby
		cmedia.comments = mediastate[0].comments
		cmedia.assignedto = commembership[0].user.username
		cmedia.assignedon = commembership[0].assignedon
		currentday = datetime.datetime.now().date() 
		changedon = mediastate[0].changedon.date()
		diff = (currentday - changedon).days
		if cmedia.media.state.name == 'submitted' or cmedia.media.state.name == 'reviewStarted' or cmedia.media.state.name == 'sentToModify':
			if int(diff) > int(settings.DEADLINE):
				cmedia.deadlinepassed = True
			else:
				cmedia.deadlinepassed = False
		else:
			cmedia.deadlinepassed = False
	lstContent = list(commarticles) + list(commmedia)
	return lstContent

def view_all_content(request, pk1, state):
	u = User.objects.get(username=request.user)
	if u.groups.filter(name='curator').exists():
		community = Community.objects.get(pk=pk1)
		introduction = CommunityArticles.objects.filter(community__name='Introduction', community__parent=community, article__state__name=state)
		architecture = CommunityArticles.objects.filter(community__name='Architecture', community__parent=community, article__state__name=state)
		rituals = CommunityArticles.objects.filter(community__name='Rituals', community__parent=community, article__state__name=state)
		ceremonies = CommunityArticles.objects.filter(community__name='Ceremonies', community__parent=community, article__state__name=state)
		tales = CommunityArticles.objects.filter(community__name='Tales', community__parent=community, article__state__name=state)
		moreinfo = CommunityArticles.objects.filter(community__name='More Information', community__parent=community, article__state__name=state)
		curatorrole = Roles.objects.get(name='curator')
		curatorname = CommunityMembership.objects.filter(community=community, role=curatorrole).order_by('-assignedon')[:1]
		curatorname = curatorname[0].user

		allowMerging = True
		pedingarticles = CommunityArticles.objects.filter(community__parent=community)
		pedingarticles = pedingarticles.filter( Q(article__state__name='submitted') | Q(article__state__name='reviewStarted') | Q(article__state__name='sentToModify') ).count()
		pedingmedia = CommunityMedia.objects.filter(community__parent=community)
		pedingmedia = pedingmedia.filter( Q(media__state__name='submitted') | Q(media__state__name='reviewStarted') | Q(media__state__name='sentToModify') ).count()
		if pedingarticles > 0 or pedingmedia > 0:
			allowMerging = False

		if MergedArticles.objects.filter(community=community).exists():
			merged = True
		else:
			merged = False
		media = CommunityMedia.objects.filter(community__parent=community, media__state__name=state)
		return render(request, 'view_all_content.html',{'community':community, 'introduction':introduction, 'architecture':architecture, 'rituals':rituals, 'ceremonies':ceremonies, 'tales':tales, 'moreinfo':moreinfo, 'merged':merged, 'curatorname':curatorname, 'media':media, 'allowMerging':allowMerging})
	return redirect('login')

def merge_content(request, pk):
	u = User.objects.get(username=request.user)
	if u.groups.filter(name='curator').exists():
		state = 'accepted'
		introduction = ''
		architecture = ''
		rituals = ''
		ceremonies = ''
		tales = ''
		moreinfo = ''
		articlereferences = ''
		originals = []
		originalmedia = []
		community = Community.objects.get(pk=pk)

		introductionQuery = CommunityArticles.objects.filter(community__name='Introduction', community__parent=community, article__state__name=state)
		articlereferences = '<h4><b>Introduction</b></h4>'
		for obj in introductionQuery:
			introduction += obj.article.body
			introduction += '<br />'
			articlereferences += obj.article.references
			originals.append(obj.article.pk)

		architectureQuery = CommunityArticles.objects.filter(community__name='Architecture', community__parent=community, article__state__name=state)
		articlereferences += '<h4><b>Architecture</b></h4>'
		for obj in architectureQuery:
			architecture += obj.article.body
			architecture += '<br />'
			articlereferences += obj.article.references
			originals.append(obj.article.pk)

		ritualsQuery = CommunityArticles.objects.filter(community__name='Rituals', community__parent=community, article__state__name=state)
		articlereferences += '<h4><b>Rituals</b></h4>'
		for obj in ritualsQuery:
			rituals += obj.article.body
			rituals += '<br />'
			articlereferences += obj.article.references
			originals.append(obj.article.pk)
		
		ceremoniesQuery = CommunityArticles.objects.filter(community__name='Ceremonies', community__parent=community, article__state__name=state)
		articlereferences += '<h4><b>Ceremonies</b></h4>'
		for obj in ceremoniesQuery:
			ceremonies += obj.article.body
			ceremonies += '<br />'
			articlereferences += obj.article.references
			originals.append(obj.article.pk)
		
		talesQuery = CommunityArticles.objects.filter(community__name='Tales', community__parent=community, article__state__name=state)
		articlereferences += '<h4><b>Tales</b></h4>'
		for obj in talesQuery:
			tales += obj.article.body
			tales += '<br />'
			articlereferences += obj.article.references
			originals.append(obj.article.pk)
		
		moreinfoQuery = CommunityArticles.objects.filter(community__name='More Information', community__parent=community, article__state__name=state)
		articlereferences += '<h4><b>More Information</b></h4>'
		for obj in moreinfoQuery:
			moreinfo += obj.article.body
			moreinfo += '<br />'
			articlereferences += obj.article.references
			originals.append(obj.article.pk)

		medias = CommunityMedia.objects.filter(community__parent=community, media__state__name=state)
		for obj in medias:
			originalmedia.append(obj.media.pk)

		state = States.objects.get(name='merged')
		merged = MergedArticles.objects.create(
			community = community,
			introduction = introduction,
			architecture = architecture,
			rituals = rituals,
			ceremonies = ceremonies,
			tales = tales,
			moreinfo = moreinfo,
			state = state,
			articlereferences = articlereferences,
			changedby = request.user,
			changedon = datetime.datetime.now(),
			originalarticles = originals,
			originalmedia = originalmedia
		)

		MergedArticleStates.objects.create (
			mergedarticle = merged,
			state = state,
			changedby = request.user,
			changedon = datetime.datetime.now(),
			introduction = introduction,
			architecture = architecture,
			rituals = rituals,
			ceremonies = ceremonies,
			tales = tales,
			moreinfo = moreinfo
		)
		community.contribution_status = 'Completed'
		community.save()
		return redirect('view_merged_content',pk=pk)
	return redirect('login')

def view_merged_content(request, pk):
	u = User.objects.get(username=request.user)
	if u.groups.filter(name='curator').exists() or u.groups.filter(name='icpapprover').exists():
		community = Community.objects.get(pk=pk)
		media = CommunityMedia.objects.filter(community__parent=community).filter(Q(media__state__name='accepted') | Q(media__state__name='publishedICP'))
		merged = MergedArticles.objects.get(community=community)
		statehistory = MergedArticleStates.objects.filter(mergedarticle=merged).order_by('-changedon')
		docuploadform = DocumentForm()
		return render(request, 'view_merged_content.html',{'merged':merged, 'media':media, 'statehistory':statehistory, 'docuploadform':docuploadform})
	return redirect('login')

def curate_merged(request):
	pk = request.POST['pk']
	merged = MergedArticles.objects.get(pk=pk)
	status = request.POST['status']
	comments = ''
	publishedlink = ''
	filepath = ''
	fileurl = ''
	to = []

	if status == 'sentForReview' or status == 'sentForApproval': 
		state = States.objects.get(name=status)
		comments = request.POST['reason']
		if status == 'sentForReview':
			users = User.objects.filter(groups__name='icpreviewer').exclude(username='admin').values_list('email', flat=True)
		if status == 'sentForApproval':
			users = User.objects.filter(groups__name='icpapprover').exclude(username='admin').values_list('email', flat=True)
		for user in users:
			to.append(user)
		filepath, fileurl = convert_to_docx(merged)
		merged.document_sent = fileurl

	if status == 'recurate':
		state = States.objects.get(name='merged')
		comments = request.POST['reason']
		originalstate = States.objects.get(name='accepted')
		change_state_orginal_contributions(merged, originalstate, request)

		reviewers = User.objects.filter(groups__name='icpreviewer').exclude(username='admin').values_list('email', flat=True)
		approvers = User.objects.filter(groups__name='icpapprover').exclude(username='admin').values_list('email', flat=True)
		users = []
		users = list(reviewers) + list(approvers)
		for user in users:
			to.append(user)

	# if status == 'accept':
	# 	state = States.objects.get(name='publish')
	# 	change_state_orginal_contributions(merged, state, request)
	# 	to.append(merged.changedby.email)

	if status == 'publishedonicp':
		state = States.objects.get(name='publishedICP')
		publishedlink = request.POST['publishedlink']
		merged.publishedlink = publishedlink
		set_published_link(merged, publishedlink)
		change_state_orginal_contributions(merged, state, request)  
		to.append(merged.changedby.email)

		articleids = merged.originalarticles
		mediaids = merged.originalmedia
		articleids = ast.literal_eval(articleids)
		mediaids = ast.literal_eval(mediaids)
		articles = Articles.objects.filter(pk__in=articleids).values_list('created_by__email', flat=True).distinct()
		medias = Media.objects.filter(pk__in=mediaids).values_list('created_by__email', flat=True).distinct()
		to = list(articles) + list(medias)
		to = list(set(to))

	merged.state = state
	merged.save()

	MergedArticleStates.objects.create (
		mergedarticle = merged,
		state = state,
		changedby = request.user,
		changedon = datetime.datetime.now(),
		introduction = merged.introduction,
		architecture = merged.architecture,
		rituals = merged.rituals,
		ceremonies = merged.ceremonies,
		tales = merged.tales,
		moreinfo = merged.moreinfo,
		comments = comments,
		document_sent = fileurl
	)
	sendEmail_merged_content_curated(to, request.user.email, status, merged.community.name, comments, request.META.get('HTTP_REFERER'), publishedlink, filepath)
	return redirect('view_merged_content',pk=merged.community.pk)

def convert_to_docx(merged):

	document = Document()
	new_parser = HtmlToDocx()

	document.add_heading(merged.community.name, 0)

	document.add_heading('Introduction', level=1)
	new_parser.add_html_to_document(merged.introduction, document)

	document.add_heading('Architecture', level=1)
	new_parser.add_html_to_document(merged.architecture, document)

	document.add_heading('Rituals', level=1)
	new_parser.add_html_to_document(merged.rituals, document)

	document.add_heading('Ceremonies', level=1)
	new_parser.add_html_to_document(merged.ceremonies, document)

	document.add_heading('Tales', level=1)
	new_parser.add_html_to_document(merged.tales, document)

	document.add_heading('More information', level=1)
	new_parser.add_html_to_document(merged.moreinfo, document)

	document.add_heading('References', level=1)
	new_parser.add_html_to_document(merged.articlereferences, document)

	document.add_heading('Gallery', level=1)

	mediaids = merged.originalmedia
	mediaids = ast.literal_eval(mediaids)
	medias = Media.objects.filter(pk__in=mediaids)
	current_site = Site.objects.get_current()
	domain = current_site.domain
	for media in medias:
		if media.mediafile:
			name = domain + settings.MEDIA_URL + f'{media.mediafile}'
			document.add_paragraph(name, style='List Bullet')
		else:
			name = media.medialink
			document.add_paragraph(name, style='List Bullet')
		document.add_paragraph('References', style='List Bullet')
		new_parser.add_html_to_document(media.references, document)
		new_parser.add_html_to_document('<br />', document)


	path = settings.MEDIA_ROOT + "/writeup/" + f'{merged.community.pk}' 
	try:
		os.makedirs(path, exist_ok = True)
	except OSError as error:
		print("Directory can not be created")

	now = datetime.datetime.now()
	filename = f'{now.year}' + "_" + f'{now.month}' + "_" + f'{now.day}' + "_" + f'{now.hour}' + f'{now.minute}' + "_" + f'{merged.community.pk}' + "_" + merged.community.name + ".docx"
	filepath = settings.MEDIA_ROOT + "/writeup/" + f'{merged.community.pk}' + "/" + filename
	fileurl = "writeup/" + f'{merged.community.pk}' + "/" + filename
	document.save(filepath)

	return filepath, fileurl

def set_published_link(merged, publishedlink):
	originalarticles = merged.originalarticles
	originalarticles = ast.literal_eval(originalarticles)
	articles = Articles.objects.filter(pk__in=originalarticles)
	for article in articles:
		article.publishedlink = publishedlink
		article.save()
	originalmedia = merged.originalmedia
	originalmedia = ast.literal_eval(originalmedia)
	medias = Media.objects.filter(pk__in=originalmedia)
	for media in medias:
		media.publishedlink = publishedlink
		media.save()


def change_state_orginal_contributions(merged, state, request):
	originalarticles = merged.originalarticles
	originalarticles = ast.literal_eval(originalarticles)
	articles = Articles.objects.filter(pk__in=originalarticles)
	for article in articles:
		article.state = state
		article.save()
		ArticleStates.objects.create(
			article = article,
			state = state,
			changedby = request.user,
			changedon = datetime.datetime.now(),
			body = article.body,
		)

	originalmedia = merged.originalmedia
	originalmedia = ast.literal_eval(originalmedia)
	medias = Media.objects.filter(pk__in=originalmedia)
	for media in medias:
		media.state = state
		media.save()
		MediaStates.objects.create(
			media = media,
			state = state,
			changedby = request.user,
			changedon = datetime.datetime.now()
		)

def get_districts(request):
	state = request.GET.get('state')
	results = Locations.objects.filter(state=state).values_list('district', flat=True).order_by('district').distinct()
	return render(request, 'address_dropdown_list_options.html', {'results': results})

def get_cities(request):
	district = request.GET.get('district')
	if district == 'Others':
		state = request.GET.get('state')
		results = Locations.objects.filter(state=state).values_list('city', flat=True).order_by('city').distinct()
	else:
		results = Locations.objects.filter(district=district).values_list('city', flat=True).order_by('city').distinct()
	return render(request, 'address_dropdown_list_options.html', {'results': results})

def get_pincodes(request):
	city = request.GET.get('city')
	if city == 'Others':
		state = request.GET.get('state')
		results = Locations.objects.filter(state=state).values_list('pincode', flat=True).order_by('pincode').distinct()
	else:
		results = Locations.objects.filter(city=city).values_list('pincode', flat=True).order_by('pincode').distinct()
	return render(request, 'address_dropdown_list_options.html', {'results': results})

def upload_received_document(request):
	pk = request.POST['pk']
	mpk = request.POST['mpk']
	commpk = request.POST['commpk']
	form = DocumentForm(request.POST, request.FILES)
	if form.is_valid():
		merged = MergedArticles.objects.get(pk=pk)
		merged.document_received = request.FILES['docfile']
		merged.save()

		mergedarticlestate = MergedArticleStates.objects.get(pk=mpk)
		mergedarticlestate.document_received = merged.document_received
		mergedarticlestate.save()

		return redirect(view_merged_content, commpk)
